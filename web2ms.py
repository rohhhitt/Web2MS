import requests
from bs4 import BeautifulSoup
import csv
import os
import time
from docx import Document
from docx.shared import Pt

# Configurations
RETRY_LIMIT = 2
OUTPUT_FOLDER = "scraped_docs"
HEADERS = {"User-Agent": "Mozilla/5.0"}

# Ensure output folder exists
os.makedirs(OUTPUT_FOLDER, exist_ok=True)

def get_page_content(url):
    """Fetches page content with retry logic."""
    for attempt in range(RETRY_LIMIT):
        try:
            response = requests.get(url.strip(), headers=HEADERS, timeout=10)
            response.raise_for_status()
            return response.text
        except requests.RequestException as e:
            print(f"⚠️ Attempt {attempt + 1} failed for {url}: {e}")
            time.sleep(2)
    return None

def extract_transcripts(soup):
    """Extracts transcripts from embedded videos if available."""
    transcripts = []
    video_tags = soup.find_all("video")
    for video in video_tags:
        transcript_tag = video.find_next("track", kind="subtitles")
        if transcript_tag and transcript_tag.has_attr("src"):
            transcripts.append(transcript_tag["src"])
    return transcripts

def extract_content(url, identifier):
    """Extracts main content while maintaining relative font sizes."""
    html = get_page_content(url)
    if not html:
        print(f"❌ Skipping {url} due to repeated failures.")
        return
    
    soup = BeautifulSoup(html, "html.parser")
    content = soup.find("main") or soup.body
    if not content:
        print(f"❌ No content found for {url}")
        return
    
    doc = Document()
    doc.add_heading(f"Extracted Content from {url}", level=1)
    
    for element in content.find_all(["p", "h1", "h2", "h3", "h4", "h5", "h6"]):
        text = element.get_text(strip=True)
        if not text:
            continue
        
        para = doc.add_paragraph(text)
        font = para.runs[0].font
        
        # Adjust font size based on tag
        font_sizes = {"h1": 24, "h2": 20, "h3": 18, "h4": 16, "h5": 14}
        font.size = Pt(font_sizes.get(element.name, 12))
    
    # Extract transcripts if available
    transcripts = extract_transcripts(soup)
    if transcripts:
        doc.add_page_break()
        doc.add_heading("Video Transcripts", level=2)
        for transcript in transcripts:
            doc.add_paragraph(transcript)
    
    output_path = os.path.join(OUTPUT_FOLDER, f"{identifier}.docx")
    doc.save(output_path)
    print(f"✅ Saved: {output_path}")

def detect_delimiter(csv_file):
    """Detects delimiter (`,` or `;`) in the CSV file."""
    with open(csv_file, newline="", encoding="utf-8") as file:
        sample = file.read(1024)
        return ";" if sample.count(";") > sample.count(",") else ","

def clean_row(row):
    """Cleans up CSV rows by removing extra semicolons and empty values."""
    row = [col.strip().strip(";") for col in row if col.strip()]
    return row if len(row) == 2 and row[0].startswith("http") else None

def process_csv(csv_file):
    """Reads CSV and processes URLs."""
    delimiter = detect_delimiter(csv_file)
    
    with open(csv_file, newline="", encoding="utf-8") as file:
        reader = csv.reader(file, delimiter=delimiter)
        next(reader, None)  # Skip header row if present

        for row in reader:
            cleaned_row = clean_row(row)
            if not cleaned_row:
                print(f"⚠️ Skipping invalid row: {row}")
                continue
            
            url, identifier = cleaned_row
            extract_content(url, identifier)
    
    print("\n🎉 Script Execution Completed!")
    print(f"📂 Word files saved in: {os.path.abspath(OUTPUT_FOLDER)}")

# Run script
process_csv("urls.csv")
